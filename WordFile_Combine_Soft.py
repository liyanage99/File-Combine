import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

def combine_files():
    directory = entry_directory.get()
    if not directory:
        messagebox.showerror("Error", "Please select a directory.")
        return

    output_directory = filedialog.askdirectory()
    if not output_directory:
        return

    for filename in os.listdir(directory):
        if filename.endswith("-a.docx"):
            file_a_path = os.path.join(directory, filename)
            file_j_path = os.path.join(directory, filename.replace("-a.docx", "-J.docx"))
            if os.path.exists(file_j_path):
                output_path = os.path.join(output_directory, filename.replace("-a.docx", "-combined.docx"))
                combine_word_files(file_a_path, file_j_path, output_path)
                messagebox.showinfo("Success", "Files combined successfully.")
            else:
                messagebox.showerror("Error", f"No matching file found for {filename}.")

def combine_word_files(file1, file2, output_file):
    doc1 = Document(file1)
    doc2 = Document(file2)

    combined_doc = Document()

    for element in doc1.element.body:
        combined_doc.element.body.append(element)

    combined_doc.add_page_break()

    for element in doc2.element.body:
        combined_doc.element.body.append(element)

    combined_doc.save(output_file)

# Create the main window
root = tk.Tk()
root.title("Combine Word Files")

# Create entry field for the directory
entry_directory = tk.Entry(root, width=50)
entry_directory.grid(row=0, column=0, padx=5, pady=5)

# Create button for selecting the directory
btn_select_directory = tk.Button(root, text="Select Directory", command=lambda: select_directory(entry_directory))
btn_select_directory.grid(row=0, column=1, padx=5, pady=5)

# Create button for combining files
btn_combine = tk.Button(root, text="Combine Files", command=combine_files)
btn_combine.grid(row=1, column=0, columnspan=2, padx=5, pady=5)

def select_directory(entry):
    directory = filedialog.askdirectory()
    entry.delete(0, tk.END)
    entry.insert(0, directory)

# Run the GUI
root.mainloop()
